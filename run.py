import pandas as pd

df = pd.read_csv('Membership_Track.csv')
# df = pd.read_excel('Membership_Track.xlsx', sheet_name="members")

from fpdf import FPDF
from datetime import date

today = date.today()
print("Today's date:", today)

class PDF(FPDF):
    def header(self):
        # Logo
        self.image('logo.png', 88, 8, 43)
        # Arial bold 15
        self.set_font('Arial', 'B', 15)
        # Move to the right
        self.cell(80)
        # Title
        #self.cell(30, 10, 'Title', 1, 0, 'C')
        # Line break
        self.set_draw_color(0, 80, 180)
        self.set_line_width(10)
        self.ln(20)

    # Page footer
    def footer(self):
        # Position at 1.5 cm from bottom
        self.set_y(-15)
        # Arial italic 8
        self.set_font('Courier', 'B', 8)
        # Page number
        self.set_text_color(0, 80, 180)
        self.set_font('GillSansMT', '', 10)
        self.cell(0, 10, "https://jadhavsujit4.github.io/blogging", 0, 0, 'C',False, "https://jadhavsujit4.github.io/blogging")

# Instantiation of inherited class
# len(df)
for i in range(len(df)):
    print("Creating pdf number ", i+1)
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()

    pdf.add_font('GillSansMT', '', 'Gill Sans MT Medium.ttf', uni=True)
    pdf.add_font('GillSansMTBold', '', 'Gill Sans MT Bold.ttf', uni=True)

    pdf.set_font('GillSansMT', '', 10)

    pdf.set_draw_color(0, 80, 180)
    pdf.set_line_width(0.6)
    pdf.line(0, 28, 300, 28)

    pdf.cell(10, 10, 'Date: ' + str(today.strftime('%d-%m-%Y')), 0, 1)

    pdf.ln(h = 2)

    pdf.set_font('GillSansMT', '', 10)
    pdf.set_text_color(50, 135, 220)
    pdf.cell(10, 8, df.iloc[i]["Name of Business Owner"], 0, 1)
    pdf.cell(10, 8, df.iloc[i]["Name of Enterprise"], 0, 1)
    pdf.cell(10, 8, df.iloc[i]["City"], 0, 1)

    pdf.set_text_color(50, 135, 220)
    pdf.cell(0, 10, "Your Membership No: " + str(df.iloc[i]["Membership No."]), 0, 1,"C")
    pdf.ln(h = '')

    pdf.set_text_color(0, 0, 0)
    # pdf.cell(10, 10, "Dear Mr. " + str(df.iloc[i]["Name of Business Owner"]) + ",", 0, 1, "U")
    # When gender column is added
    salutation = ""
    if str(df.iloc[i]["Gender"]) == "nan" :
        salutation = ""
    else:
        salutation = str(df.iloc[i]["Gender"]) + " "
    pdf.cell(0, 10, "Dear " + str(salutation) + str(df.iloc[i]["Name of Business Owner"]) + ",", 0, 1, "U")

    pdf.set_font('GillSansMTBold', '', 10)
    pdf.cell(10, 10, "We are delighted that you are now a member of Our some Network which will help you!", 0, 1,"B")
    pdf.ln(h=4)

    pdf.set_font('GillSansMT', '', 10)
    pdf.multi_cell(0, 5, "On behalf of the Management team of Some enterprise that we don't have name for, I wish you a warm welcome and thank you for choosing to grow your business with us.", 0, 1)
    pdf.ln(h=4)
    
    if str(df.iloc[i]["International"]) == "yes" :
        pdf.multi_cell(0, 5, "Some enterprise that we don't have name for equips business owners with necessary knowledge, mentoring, resources, tools and consulting by industry experts to grow their businesses exponentially.")
    else:
        pdf.multi_cell(0, 5, "Some enterprise that we don't have name for provides a structured and methodical approach to dealing with critical business challenges, guided by industry experts in specialised domains, through mentoring of business owners.")
    pdf.ln(h=4)

    pdf.multi_cell(0, 5, "Your membership automatically includes access to various network platforms with orgs and industry experts from varying business domains. We have several exciting engagement opportunities including Webinars, Accelerator Cohort Programs, Mentoring, Consulting, Networking with Experts and Stalwarts. Our flagship service – Accelerator Cohort Programs are time-bound professional engagement programs, that are curated by experts and result oriented for individual businesses. I would urge you to review and consider them for your business needs.")
    pdf.ln(h=4)

    pdf.multi_cell(0, 5, "As much as we will be happy to keep you abreast of Industry updates, we shall also be happy to help you find what you are looking for. Feel free to write to our support team. We send out most communications via email, so please make sure to keep your contact information updated with us.")
    pdf.ln(h=4)

    pdf.multi_cell(0, 5, "Start taking advantage of all your Member Benefits –")
    pdf.image('bullets.png', 17, 190.5, 6.5, 21, "PNG")
    pdf.cell(10, 5, "              4 Free webinars by experts on key business focus areas", 0, 1)
    pdf.cell(30, 5, "              1 Networking event to connect with co-members and mentors", 0, 1)
    pdf.cell(30, 5, "              Access to marquee events at nominal price", 0, 1)
    pdf.cell(30, 5, "              Discounts on Accelerator Cohort programs", 0, 1)
    pdf.ln(h=4)

    pdf.multi_cell(0, 5, "Your membership is free under the Covid19 relief scheme and valid till 31st May, 2021. Keep visiting our website to stay connected, and again, welcome! You have made a sound investment for the future of your enterprise!")
    pdf.ln(h=4)

    pdf.cell(0, 5, "Warm Regards,")
    pdf.image('image.png', 15, 235, 15, 15, "PNG")
    pdf.ln(h=4)

    pdf.set_font('GillSansMTBold', '', 10)
    pdf.cell(0, 35, "SUJIT JADHAV")
    pdf.ln(h=20)

    pdf.set_font('GillSansMT', '', 10)
    pdf.cell(0, 8, "Chief Executive Officer", 0, 1)
    # pdf.ln(h=28)
    pdf.set_text_color(50, 135, 220)
    pdf.cell(0, 4, "Some enterprise that we don't have name for")

    pdf.set_text_color(0, 0, 0)
    pdf.ln(h=4)
    pdf.cell(30, 8, "Write to me at: jadhav.sujit4@gmail.com")
    # pdf.ln(h=4)

    pdf.line(0, 275, 300, 275)
    print(df.iloc[i]["Membership No."])
    pdfName = str("Membership Welcome Letter " + str(df.iloc[i]["Membership No."]) + ".pdf")
    pdf.output(pdfName, 'F')
    

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Emails to be pasted', 0)

for i in range(len(df)):
    p = document.add_paragraph('')
    salutation = ""
    if str(df.iloc[i]["Gender"]) == "nan" :
        salutation = ""
    else:
        salutation = str(df.iloc[i]["Gender"]) + " "
    p.add_run('Dear ' + str(salutation) + str(df.iloc[i]["Name of Business Owner"]) + ',').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    document.add_paragraph('Congratulations on your membership with us!')
    r = document.add_paragraph('Please find the attached Welcome kit comprising of your Membership Letter and our Brochure. \n')
    r.add_run('Please download Zoom application for some of our webinar engagements till we have to maintain social distancing norms.')
    q = document.add_paragraph('')
    q.add_run('Please respond to this mail confirming your membership.').bold = True
    document.add_paragraph('For any queries related to membership or any of our services, please feel free to write back to us!')
    document.add_paragraph('Stay safe and keep your family safe!')
    document.add_paragraph('Best Wishes!')
    document.add_paragraph('***************************************************************************')

document.save('Emails.docx')
